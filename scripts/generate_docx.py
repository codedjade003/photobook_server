#!/usr/bin/env python3
"""Generate the PhotoBook Backend Technical Summary & Sprint Checklist as a .docx file."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ── Page Setup ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)

# Heading styles
for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    if level == 1:
        h.font.size = Pt(22)
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(12)
    elif level == 2:
        h.font.size = Pt(16)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
    else:
        h.font.size = Pt(13)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

# ── Helper functions ───────────────────────────────────────

def add_colored_paragraph(doc, text, color=None, bold=False, size=None, alignment=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    if alignment is not None:
        p.alignment = alignment
    return p

def add_checkbox(doc, text, checked=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5)
    box = "☑" if checked else "☐"
    run = p.add_run(f"{box}  {text}")
    run.font.size = Pt(10.5)
    if color:
        run.font.color.rgb = color
    return p

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_section_heading(doc, number, title, color=None):
    h = doc.add_heading(f"{number}. {title}", level=2)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

# ═══════════════════════════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

add_colored_paragraph(
    doc, "PHOTOBOOK",
    color=RGBColor(0x1a, 0x1a, 0x2e),
    bold=True,
    size=Pt(36),
    alignment=WD_ALIGN_PARAGRAPH.CENTER
)

add_colored_paragraph(
    doc, "Backend Technical Summary",
    color=RGBColor(0x55, 0x55, 0x77),
    bold=False,
    size=Pt(18),
    alignment=WD_ALIGN_PARAGRAPH.CENTER
)

add_colored_paragraph(
    doc, "& Sprint Checklist",
    color=RGBColor(0x55, 0x55, 0x77),
    bold=False,
    size=Pt(18),
    alignment=WD_ALIGN_PARAGRAPH.CENTER
)

doc.add_paragraph()
add_horizontal_rule(doc)
doc.add_paragraph()

add_colored_paragraph(
    doc, f"Date: {datetime.date.today().strftime('%B %d, %Y')}",
    color=RGBColor(0x77, 0x77, 0x99),
    size=Pt(12),
    alignment=WD_ALIGN_PARAGRAPH.CENTER
)
add_colored_paragraph(
    doc, "Audience: Full Team (Technical & Non-Technical)",
    color=RGBColor(0x77, 0x77, 0x99),
    size=Pt(12),
    alignment=WD_ALIGN_PARAGRAPH.CENTER
)
add_colored_paragraph(
    doc, "Goal: Ship Stable Backend to Staging in ≤ 1 Week",
    color=RGBColor(0x77, 0x77, 0x99),
    size=Pt(12),
    alignment=WD_ALIGN_PARAGRAPH.CENTER
)
add_colored_paragraph(
    doc, "Classification: Internal — Engineering Team",
    color=RGBColor(0x99, 0x99, 0xBB),
    size=Pt(10),
    alignment=WD_ALIGN_PARAGRAPH.CENTER
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════

doc.add_heading("Table of Contents", level=1)

toc_items = [
    ("1.", "What We Built (In Plain English)", "3"),
    ("2.", "Overall Health Check", "4"),
    ("3.", "Master Checklist Overview", "5"),
    ("4.", "Backend Engineer Checklist — Critical Fixes", "6"),
    ("5.", "Backend Engineer Checklist — High Priority", "7"),
    ("6.", "Backend Engineer Checklist — Medium Priority", "8"),
    ("7.", "Frontend Engineer Test Checklist", "9"),
    ("8.", "Estimated 1-Week Timeline", "12"),
    ("9.", "Risks & Dependencies", "13"),
    ("10.", "Definition of 'Staging Ready'", "13"),
    ("11.", "Quick Reference", "14"),
]

for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"{num}  {title}")
    run.font.size = Pt(11)
    tab_run = p.add_run(f"  ··········  {page}")
    tab_run.font.size = Pt(10)
    tab_run.font.color.rgb = RGBColor(0x99, 0x99, 0xBB)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  SECTION 1: WHAT WE BUILT
# ═══════════════════════════════════════════════════════════════

doc.add_heading("1. What We Built (In Plain English)", level=1)

doc.add_paragraph(
    "PhotoBook is a marketplace that connects people who need photographers with "
    "professional photographers. Think of it like Upwork or Fiverr, but specifically "
    "for photography in Nigeria."
)

doc.add_paragraph("What the backend does today:", style='List Bullet')

features = [
    ("Sign up & Login", "People can create accounts with email/password or Google, and log in."),
    ("Profiles", "Photographers set up their business name, bio, tags; clients set their location."),
    ("Portfolio", "Photographers upload photos & videos stored in the cloud (Backblaze B2)."),
    ("Rate Cards", "Photographers list their prices for different services."),
    ("Booking", "Clients can book a session — pick event type, date, time, and location."),
    ("Search & Discovery", "Search photographers by name, location, tags, rating, and more. Includes smart similarity matching and trending tags."),
    ("Real-time Messaging", "Clients and photographers chat in real-time with encrypted messages via Socket.io."),
    ("Video Call Plumbing", "WebRTC signaling (offer/answer) is built for future video calls."),
    ("API Documentation", "Every endpoint is documented at /api-docs (Swagger) with 'Try it out' buttons."),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "Feature"
hdr[1].text = "What It Means"
for f, d in features:
    row = table.add_row()
    row.cells[0].text = f
    row.cells[1].text = d

doc.add_paragraph()
doc.add_paragraph("What's NOT built yet (planned for future sprints):", style='List Bullet')

missing = [
    "Payments — no way to actually pay for a booking yet.",
    "Reviews — the database table exists but no API endpoints are built.",
    "Booking lifecycle — sessions are created but can't be confirmed, completed, or cancelled.",
    "Admin dashboard — no tools for managing users or content.",
    "SMS notifications — config stub exists but not implemented.",
]
for m in missing:
    p = doc.add_paragraph(m, style='List Bullet')
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xCC, 0x55, 0x00)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  SECTION 2: HEALTH CHECK
# ═══════════════════════════════════════════════════════════════

doc.add_heading("2. Overall Health Check", level=1)

health_data = [
    ("Code Structure", "⭐⭐⭐⭐  Good", "Well-organized, easy to navigate. Clean 3-layer architecture."),
    ("Security", "⭐⭐  Needs Work", "Critical fixes needed before any public exposure."),
    ("Performance", "⭐⭐⭐  OK", "Fine for current scale; needs pagination before growth."),
    ("Completeness", "⭐⭐⭐  ~60%", "Core features work; payments, reviews, booking lifecycle missing."),
    ("Documentation", "⭐⭐⭐⭐  Good", "Swagger docs are thorough with examples for every endpoint."),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = "Area"
hdr[1].text = "Rating"
hdr[2].text = "Notes"
for area, rating, notes in health_data:
    row = table.add_row()
    row.cells[0].text = area
    row.cells[1].text = rating
    row.cells[2].text = notes

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  SECTION 3: MASTER CHECKLIST OVERVIEW
# ═══════════════════════════════════════════════════════════════

doc.add_heading("3. Master Checklist — Priority Legend", level=1)

legend = [
    ("🔴  CRITICAL", "Must fix before staging deployment. Security or crash-risk issues.", RGBColor(0xCC, 0x00, 0x00)),
    ("🟠  HIGH PRIORITY", "Should fix during this sprint. Code quality, stability, and cleanup.", RGBColor(0xCC, 0x66, 0x00)),
    ("🟡  MEDIUM PRIORITY", "Can follow up next sprint. Developer experience improvements.", RGBColor(0xBB, 0x99, 0x00)),
    ("🟢  NICE TO HAVE", "Future sprints. New features and major improvements.", RGBColor(0x33, 0x99, 0x33)),
    ("☑  DONE", "Already completed and verified.", RGBColor(0x33, 0x99, 0x33)),
]

for label, desc, color in legend:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}  —  {desc}")
    run.font.color.rgb = color
    run.bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  SECTION 4: BACKEND — CRITICAL
# ═══════════════════════════════════════════════════════════════

add_section_heading(doc, 4, "Backend Engineer Checklist — 🔴 CRITICAL Fixes", RGBColor(0xCC, 0x00, 0x00))
doc.add_paragraph("These MUST be completed before the staging deployment. They address security vulnerabilities and server crashes.")

critical_items = [
    ("Rotate ALL credentials",
     "The .env file contains real passwords, API keys, and secrets. Change: Supabase DB password, "
     "Redis password, Resend API key, Google OAuth secret, JWT secret, B2 application key, Session secret.\n"
     "⚠️  Risk if skipped: VERY HIGH — account compromise.\n"
     "⏱  Effort: ~2 hours",
     False),
    ("Replace Math.random() with crypto.randomInt() for security codes",
     "In src/services/auth.service.js, the generateCode() function uses Math.random() for "
     "6-digit verification codes and password reset codes — this is predictable.\n"
     "⚠️  Risk if skipped: HIGH — codes can be guessed by attackers.\n"
     "⏱  Effort: ~5 minutes",
     False),
    ("Wire 2FA into the login flow",
     "2FA can be set up (POST /api/auth/2fa/setup & /confirm work) but it is NEVER checked "
     "during login. Add a check in loginUser() inside src/services/auth.service.js.\n"
     "⚠️  Risk if skipped: MEDIUM — 2FA is currently decorative.\n"
     "⏱  Effort: ~2 hours",
     False),
    ("Fix crash on portfolio delete with dev override password",
     "In src/controllers/portfolio.controller.js, the deletePortfolioItemController "
     "accesses req.user.id but when using the dev override password (no token), "
     "req.user is undefined — this crashes the server.\n"
     "⚠️  Risk if skipped: HIGH — server crash.\n"
     "⏱  Effort: ~15 minutes",
     False),
    ("Add request body size limit",
     "express.json() currently has NO limit. Add: app.use(express.json({ limit: '10mb' })) in src/app.js.\n"
     "⚠️  Risk if skipped: MEDIUM — memory exhaustion attack.\n"
     "⏱  Effort: ~2 minutes",
     False),
    ("Add Helmet.js security headers",
     "Install helmet (npm install helmet), add app.use(helmet()) to src/app.js.\n"
     "Adds CSP, X-Frame-Options, X-Content-Type-Options, and other standard security headers.\n"
     "⚠️  Risk if skipped: MEDIUM — missing standard protections.\n"
     "⏱  Effort: ~10 minutes",
     False),
    ("Configure CORS properly",
     "Currently cors() allows ANY website to call the API. Restrict to your frontend URL(s).\n"
     "⚠️  Risk if skipped: LOW-MEDIUM.\n"
     "⏱  Effort: ~5 minutes",
     False),
]

for title, desc, checked in critical_items:
    p = doc.add_paragraph()
    run = p.add_run(f"☐  {title}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1)
    r2 = p2.add_run(desc)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  SECTION 5: BACKEND — HIGH PRIORITY
# ═══════════════════════════════════════════════════════════════

add_section_heading(doc, 5, "Backend Engineer Checklist — 🟠 HIGH PRIORITY", RGBColor(0xCC, 0x66, 0x00))
doc.add_paragraph("These should be completed during this sprint for code quality and stability.")

high_items = [
    ("Remove dead/unsafe files",
     "Delete: Creatives token.txt (contains old JWT tokens — security risk), "
     "template.txt & template2.txt (replaced by email.templates.js), "
     "8 placeholder route files (health, bookings, inquiries, messages, payments, quotes, reviews, admin).\n⏱ ~10 min",
     False),
    ("Consolidate duplicate migration folders",
     "db/migrations/ and src/db/migrations/ overlap. Merge into src/db/migrations/, delete the other.\n⏱ ~30 min",
     False),
    ("Delete dead OAuth code",
     "In src/services/auth.service.js, handleGoogleOAuthCallback is unused and duplicates findOrCreateOAuthUser.\n⏱ ~5 min",
     False),
    ("Extract shared helper functions",
     "parsePositiveInt is copy-pasted in 4+ files. Move to src/utils/helpers.js. Same for parseBoolean and parseTags.\n⏱ ~1 hour",
     False),
    ("Update STRUCTURE.txt",
     "It references folders that don't exist (docs/, tests/, jobs/, models/, seeds/).\n⏱ ~15 min",
     False),
    ("Add pagination to list endpoints",
     "GET /api/portfolio/me, GET /api/sessions/me, GET /api/rate-card/me return ALL results with no limit.\n⏱ ~3 hours",
     False),
    ("Replace bcryptjs with bcrypt (native)",
     "The pure-JS version is ~2x slower. npm uninstall bcryptjs && npm install bcrypt. Same API.\n⏱ ~30 min",
     False),
    ("Replace speakeasy with otplib",
     "Speakeasy was last updated in 2019 (unmaintained). otplib is the modern, maintained alternative.\n⏱ ~1 hour",
     False),
]

for title, desc, checked in high_items:
    p = doc.add_paragraph()
    run = p.add_run(f"☐  {title}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1)
    r2 = p2.add_run(desc)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  SECTION 6: BACKEND — MEDIUM PRIORITY
# ═══════════════════════════════════════════════════════════════

add_section_heading(doc, 6, "Backend Engineer Checklist — 🟡 MEDIUM PRIORITY", RGBColor(0xBB, 0x99, 0x00))
doc.add_paragraph("Developer experience and operational improvements. Can spill into next sprint if needed.")

medium_items = [
    ("Add ESLint + Prettier — standardize code formatting. ⏱ ~1 hour", False),
    ("Add structured logging (pino) — replace console.log with production-grade JSON logs. ⏱ ~2 hours", False),
    ("Add compression middleware — gzip API responses. npm install compression. ⏱ ~10 min", False),
    ("Standardize API response shapes — some return {items}, others {sessions}, etc. Unify. ⏱ ~4 hours", False),
    ("Remove deprecated route aliases — /api/auth/resend-verification and /api/profiles/creative. ⏱ ~10 min", False),
    ("Add Docker support — Dockerfile + docker-compose.yml for consistent dev environment. ⏱ ~2 hours", False),
    ("Deploy to staging — set up Render/Railway with staging environment variables. ⏱ ~2 hours", False),
    ("Add .nvmrc file to pin Node.js version. ⏱ ~2 min", False),
]

for title, checked in medium_items:
    p = doc.add_paragraph()
    run = p.add_run(f"☐  {title}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xBB, 0x99, 0x00)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  SECTION 7: FRONTEND ENGINEER CHECKLIST
# ═══════════════════════════════════════════════════════════════

add_section_heading(doc, 7, "Frontend Engineer Test Checklist", RGBColor(0x33, 0x66, 0xAA))

doc.add_paragraph(
    "How to test: Go to the staging server's /api-docs (Swagger UI). Every endpoint has a "
    "\"Try it out\" button. You can test the entire API without writing any backend code."
)
doc.add_paragraph()

frontend_steps = [
    ("Step 1: Authentication Flow", [
        ("Sign up a test photographer", "POST /api/auth/signup  →  role: \"photographer\". Save the token."),
        ("Sign up a test client", "POST /api/auth/signup  →  role: \"client\". Save the token."),
        ("Login with both accounts", "POST /api/auth/login  →  verify tokens are returned."),
        ("Get current user", "GET /api/auth/me  →  verify user data is correct."),
        ("Update role", "PATCH /api/auth/role  →  switch between client/photographer."),
        ("Password reset flow", "POST /api/auth/password-reset/request  →  then /confirm."),
        ("Test error states", "Wrong password → 401. Unverified email → 403. Missing fields → 400."),
    ]),
    ("Step 2: Profiles", [
        ("Get my profile", "GET /api/profiles/me (as photographer, then as client)."),
        ("Update photographer profile", "PUT /api/profiles/photographer  →  business name, title, about, tags."),
        ("Update client profile", "PUT /api/profiles/client  →  location."),
        ("View public profile", "GET /api/profiles/:id  →  view another user's profile."),
    ]),
    ("Step 3: Portfolio", [
        ("Upload a photo", "POST /api/portfolio/upload  →  multipart form: file + title + tags."),
        ("Upload a video", "Same endpoint + durationSeconds. Verify video-specific validation."),
        ("List my portfolio", "GET /api/portfolio/me  →  verify signed image URLs open in browser."),
        ("Update metadata", "PATCH /api/portfolio/:itemId  →  change title, tags, isCover."),
        ("Delete an item", "DELETE /api/portfolio/:itemId  →  verify item is gone."),
    ]),
    ("Step 4: Rate Cards", [
        ("Create rate card items", "POST /api/rate-card  →  service name + pricing (fixed or contact)."),
        ("List my rate card", "GET /api/rate-card/me."),
        ("View public rate card", "GET /api/rate-card/:photographerId."),
        ("Update / Delete items", "PUT & DELETE on /api/rate-card/items/:itemId."),
    ]),
    ("Step 5: Sessions (Bookings)", [
        ("Get event types", "GET /api/sessions/event-types  →  verify dropdown list."),
        ("Book a session", "POST /api/sessions (as client, targeting photographer)."),
        ("List my sessions", "GET /api/sessions/me (as client AND as photographer — different views)."),
        ("Delete a session", "DELETE /api/sessions/:sessionId."),
    ]),
    ("Step 6: Search & Discovery", [
        ("Search photographers", "GET /api/search/users?role=photographer&q=wedding&sort=rating."),
        ("Filter by tags", "?tags=wedding,portrait&matchAllTags=true."),
        ("Filter by location", "?location=lagos."),
        ("Search portfolio", "GET /api/search/portfolio?q=bridal&mediaType=image."),
        ("Similar photographers", "GET /api/search/users/:userId/similar."),
        ("Similar portfolio items", "GET /api/search/portfolio/:itemId/similar."),
        ("Trending tags", "GET /api/search/tags/trending."),
        ("Verify pagination", "Check limit, offset, hasMore, and total in responses."),
    ]),
    ("Step 7: Messaging (REST API)", [
        ("Create a direct conversation", "POST /api/conversations  →  type: \"direct\" + participantIds."),
        ("List conversations", "GET /api/conversations."),
        ("Send a message", "POST /api/conversations/:id/messages  →  { content: \"Hello!\" }."),
        ("Fetch messages", "GET /api/conversations/:id/messages?limit=20."),
        ("Test cursor pagination", "Use nextCursor from response to fetch older messages."),
        ("Create a group conversation", "POST /api/conversations  →  type: \"group\" + 3+ participants."),
    ]),
    ("Step 8: Edge Cases & Error Handling", [
        ("Unauthorized (no token)", "Call protected endpoints → expect 401."),
        ("Wrong role (client on photographer endpoint)", "→ expect 403."),
        ("Invalid data", "Bad email, short password, missing fields → expect 400 with field errors."),
        ("Rate limiting", "Hit endpoint 100+ times rapidly → expect 429 + Retry-After header."),
        ("Duplicate email signup", "→ expect 409."),
        ("Not found", "Non-existent profile/portfolio/session → expect 404."),
    ]),
]

for step_title, items in frontend_steps:
    h = doc.add_heading(step_title, level=3)
    for sub_title, sub_desc in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(f"☐  {sub_title}")
        run.bold = True
        run.font.size = Pt(10.5)
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1.5)
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(sub_desc)
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  SECTION 8: 1-WEEK TIMELINE
# ═══════════════════════════════════════════════════════════════

add_section_heading(doc, 8, "Estimated 1-Week Timeline", RGBColor(0x33, 0x66, 0x33))

schedule = [
    ("Monday", "Backend", "Rotate credentials, fix Math.random(), fix portfolio crash, add body limit + Helmet + CORS"),
    ("Monday", "Frontend", "Auth flow testing (Step 1 — all endpoints)"),
    ("Tuesday", "Backend", "Wire 2FA into login, remove dead files, consolidate migrations, delete dead code"),
    ("Tuesday", "Frontend", "Profile + Portfolio testing (Steps 2–3)"),
    ("Wednesday", "Backend", "Extract shared helpers, add pagination, replace bcryptjs & speakeasy"),
    ("Wednesday", "Frontend", "Rate Cards + Sessions testing (Steps 4–5)"),
    ("Thursday", "Backend", "ESLint + Prettier, structured logging, compression, standardize response shapes"),
    ("Thursday", "Frontend", "Search & Discovery testing (Step 6)"),
    ("Friday", "Backend", "Docker setup, staging deployment, smoke test everything"),
    ("Friday", "Frontend", "Messaging + Edge Cases testing (Steps 7–8)"),
    ("Saturday", "Backend", "Buffer day — bug fixes from frontend findings"),
    ("Saturday", "Frontend", "Buffer day — document integration patterns for mobile app"),
    ("Sunday", "🎯 BOTH", "STAGING GO LIVE — final verification, connect mobile app"),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = "Day"
hdr[1].text = "Who"
hdr[2].text = "Tasks"
for day, who, tasks in schedule:
    row = table.add_row()
    row.cells[0].text = day
    row.cells[1].text = who
    row.cells[2].text = tasks
    if "GO LIVE" in day:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x00, 0x88, 0x00)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  SECTION 9: RISKS
# ═══════════════════════════════════════════════════════════════

add_section_heading(doc, 9, "Risks & Dependencies", RGBColor(0xCC, 0x00, 0x00))

risks = [
    ("Credential rotation breaks something",
     "HIGH — API goes down",
     "Do during low-traffic hours. Keep old credentials handy for immediate rollback."),
    ("Frontend finds API bugs late in the week",
     "MEDIUM — delays launch",
     "Backend dev stays available Saturday for hotfixes."),
    ("Payment/reviews not ready for staging",
     "LOW — not in scope",
     "Document as \"Coming Soon\" in the mobile app. Not blocking staging."),
    ("AI-generated fixes need human review",
     "LOW — quality check",
     "Every AI-assisted fix must be code-reviewed before merging."),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = "Risk"
hdr[1].text = "Impact"
hdr[2].text = "Mitigation"
for risk, impact, mitigation in risks:
    row = table.add_row()
    row.cells[0].text = risk
    row.cells[1].text = impact
    row.cells[2].text = mitigation

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
#  SECTION 10: DEFINITION OF STAGING READY
# ═══════════════════════════════════════════════════════════════

add_section_heading(doc, 10, "Definition of \"Staging Ready\"", RGBColor(0x33, 0x66, 0x33))

staging_criteria = [
    "All 7 critical security fixes applied and verified.",
    "Frontend has tested all 8 step groups and reported bugs.",
    "All reported bugs are fixed or documented as known issues.",
    "API responds correctly to all valid and invalid requests.",
    "Swagger docs match actual API behavior.",
    "Staging environment deployed and health check (/health) returns OK.",
    ".env on staging uses staging-only credentials (not production secrets).",
]

for item in staging_criteria:
    add_checkbox(doc, item, checked=False, color=RGBColor(0x33, 0x66, 0x33))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  SECTION 11: QUICK REFERENCE
# ═══════════════════════════════════════════════════════════════

add_section_heading(doc, 11, "Quick Reference — Common Commands", RGBColor(0x33, 0x33, 0x33))

commands = [
    ("Start the dev server", "npm run dev"),
    ("Run smoke test (automated API check)", "npm run test:smoke"),
    ("Generate a new message encryption key", "npm run generate:message-key"),
    ("Generate dev override password hash", "node scripts/generate-dev-hash.js \"yourPassword\""),
    ("Open Swagger API docs in browser", "open http://localhost:5001/api-docs"),
    ("Run database migrations", "psql '<DATABASE_URL>' -f src/db/migrations/001_init.sql"),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = "Action"
hdr[1].text = "Command"
for action, cmd in commands:
    row = table.add_row()
    row.cells[0].text = action
    row.cells[1].text = cmd
    # Make command monospace
    for paragraph in row.cells[1].paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)

doc.add_paragraph()
doc.add_paragraph()

# ── Footer Note ─────────────────────────────────────────────
add_horizontal_rule(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "This document was generated from a full codebase audit of 55+ source files across "
    "15 directories. For detailed technical findings, refer to the complete audit report. "
    "Generated with AI assistance."
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xBB)
run.italic = True

# ── Save ────────────────────────────────────────────────────
output_path = "/Users/mac/Desktop/pb/photobook_server/PhotoBook_Technical_Summary_Sprint_Checklist.docx"
doc.save(output_path)
print(f"✅ Document saved to: {output_path}")
